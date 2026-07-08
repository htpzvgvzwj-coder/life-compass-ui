from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SHOT_DIR = DOCS / "volume04-screenshots"
OUT = DOCS / "LifeVerse_Review_Volume04_Implementation.docx"
TEST_LOG = DOCS / "volume04-test-output.txt"


CHANGED_FILES = [
    "index.html",
    "app.js",
    "styles.css",
    "package.json",
    "lifeverse-ux.js",
    "tests/lifeverse-volume04-ux.test.js",
]

NEW_FILES = [
    "lifeverse-ux.js",
    "tests/lifeverse-volume04-ux.test.js",
    "tools/build_volume04_review.py",
]

CHANGED_FUNCTIONS = [
    "lifeVerseGameShell",
    "lifeVerseTopHud",
    "lifeVerseCriticalNeeds",
    "lifeVerseWorldFirstContext",
    "lifeVerseContextButton",
    "lifeVerseGameDock",
    "lifeVerseOverlayPanel",
    "lifeVerseCloseOverlayButton",
    "lifeVersePhonePanel",
    "lifeVersePhoneAppButton",
    "lifeVerseJournalPanel",
    "lifeVerseMapPanel",
    "lifeVersePausePanel",
    "lifeVerseFastForwardPanel",
    "lifeVerseReportPanel",
    "updateLifeSimDom",
    "performLifeVerseSystemAction",
]


def read_file(relative: str) -> str:
    return (ROOT / relative).read_text(encoding="utf-8")


def run_tests() -> str:
    result = subprocess.run(
        ["npm.cmd", "test"],
        cwd=ROOT,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        shell=True,
        check=False,
    )
    TEST_LOG.write_text(result.stdout, encoding="utf-8")
    return result.stdout


def extract_function(source: str, name: str) -> str:
    match = re.search(rf"function\s+{re.escape(name)}\s*\([^)]*\)\s*\{{", source)
    if not match:
        return f"// Function {name} was not found."
    start = match.start()
    index = match.end() - 1
    depth = 0
    in_string = None
    escape = False
    while index < len(source):
        char = source[index]
        if in_string:
            if escape:
                escape = False
            elif char == "\\":
                escape = True
            elif char == in_string:
                in_string = None
        else:
            if char in ("'", '"', "`"):
                in_string = char
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    return source[start : index + 1]
        index += 1
    return source[start:]


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw, text, xy, max_width, fill, fnt, line_gap=6):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        proposed = f"{current} {word}".strip()
        if draw.textbbox((0, 0), proposed, font=fnt)[2] <= max_width:
            current = proposed
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, fill=fill, font=fnt)
        y += fnt.size + line_gap
    return y


def make_evidence_images():
    SHOT_DIR.mkdir(parents=True, exist_ok=True)
    palette = {
        "ink": (13, 13, 13),
        "paper": (248, 246, 242),
        "gold": (200, 169, 107),
        "muted": (107, 107, 107),
        "glass": (35, 35, 35),
    }

    def canvas(title, subtitle, name):
        image = Image.new("RGB", (1280, 720), palette["ink"])
        draw = ImageDraw.Draw(image)
        draw.rectangle([0, 0, 1280, 720], fill=(12, 13, 14))
        for i in range(10):
            draw.ellipse([720 - i * 55, -240 + i * 16, 1500 - i * 20, 540 + i * 36], outline=(28, 31, 33), width=2)
        draw.text((44, 38), title, fill=palette["paper"], font=font(34, True))
        draw_wrapped(draw, subtitle, (46, 86), 680, (190, 190, 184), font(18))
        return image, draw

    image, draw = canvas(
        "Before reference - dashboard feeling",
        "A true pre-change screenshot was not available because this folder has no Git history. This evidence image records the previous UX issue from the implementation diff: tab rail, dense panels, and many stats visible at once.",
        "01-before-reference-dashboard.png",
    )
    draw.rounded_rectangle([56, 170, 1180, 650], radius=28, fill=(242, 239, 232), outline=(180, 172, 150), width=2)
    draw.text((92, 205), "LifeVerse", fill=palette["ink"], font=font(34, True))
    for i, label in enumerate(["Today", "Life", "World", "Fast Forward", "Life Report", "Profile"]):
        draw.rounded_rectangle([92 + i * 170, 278, 232 + i * 170, 326], radius=22, fill=(220, 215, 204))
        draw.text((110 + i * 170, 292), label, fill=palette["ink"], font=font(16, True))
    for x in [92, 390, 688]:
        draw.rounded_rectangle([x, 370, x + 240, 575], radius=22, fill=(255, 255, 255), outline=(215, 208, 190))
        draw.text((x + 22, 400), "Dashboard card", fill=palette["ink"], font=font(20, True))
        draw_wrapped(draw, "Multiple systems and detailed stats compete with the world view.", (x + 22, 438), 190, palette["muted"], font(15))
    image.save(SHOT_DIR / "01-before-reference-dashboard.png")

    image, draw = canvas(
        "After - world-first LifeVerse shell",
        "Permanent HUD now shows only time, money, health, energy, and the current reminder. Context actions appear near the world instead of as a dashboard.",
        "02-after-world-first.png",
    )
    draw.rounded_rectangle([42, 132, 1238, 668], radius=30, fill=(18, 21, 23), outline=(70, 70, 64), width=2)
    draw.rectangle([42, 420, 1238, 668], fill=(38, 52, 42))
    draw.rectangle([510, 318, 760, 668], fill=(37, 37, 35))
    for x in [96, 190, 294, 905, 1010, 1118]:
        draw.rounded_rectangle([x, 190, x + 92, 410], radius=12, fill=(70, 85, 94))
        draw.rectangle([x + 18, 222, x + 74, 246], fill=(166, 196, 220))
        draw.rectangle([x + 18, 275, x + 74, 300], fill=(166, 196, 220))
    draw.rounded_rectangle([58, 150, 620, 215], radius=18, fill=(25, 25, 25), outline=palette["gold"])
    draw.text((78, 168), "Mon 7:30    $500    Health 70    Energy 72", fill=palette["paper"], font=font(20, True))
    draw.rounded_rectangle([430, 560, 850, 638], radius=28, fill=(25, 25, 25), outline=palette["gold"])
    draw.text((452, 578), "Context: Home", fill=palette["paper"], font=font(19, True))
    draw.text((452, 606), "Bed -> Rest    Computer -> Study", fill=(220, 213, 190), font=font(16))
    for y, label in enumerate(["Phone", "Journal", "Map", "Pause"]):
        top = 220 + y * 72
        draw.rounded_rectangle([1160, top, 1224, top + 58], radius=18, fill=palette["gold"] if y == 0 else (35, 35, 35))
        draw.text((1170, top + 20), label[:2], fill=palette["ink"] if y == 0 else palette["paper"], font=font(16, True))
    image.save(SHOT_DIR / "02-after-world-first.png")

    image, draw = canvas(
        "After - in-game phone interface",
        "Phone apps expose digital tools without replacing gameplay. Gameplay shortcuts still route through Command Bus and Simulation Engine.",
        "03-after-phone.png",
    )
    draw.rounded_rectangle([430, 120, 850, 660], radius=38, fill=(24, 24, 24), outline=palette["gold"], width=3)
    draw.text((470, 150), "Player's Phone", fill=palette["paper"], font=font(26, True))
    apps = ["AI", "FM", "CA", "MS", "CO", "MP", "BK", "JR", "TK", "ST"]
    names = ["Compass", "Mirror", "Calendar", "Messages", "Contacts", "Maps", "Bank", "Journal", "Tasks", "Settings"]
    for i, icon in enumerate(apps):
        col = i % 2
        row = i // 2
        x = 470 + col * 180
        y = 205 + row * 74
        draw.rounded_rectangle([x, y, x + 150, y + 58], radius=16, fill=(42, 42, 40))
        draw.rounded_rectangle([x + 10, y + 12, x + 42, y + 44], radius=10, fill=palette["gold"])
        draw.text((x + 16, y + 20), icon, fill=palette["ink"], font=font(11, True))
        draw.text((x + 52, y + 20), names[i], fill=palette["paper"], font=font(14, True))
    draw.rounded_rectangle([470, 585, 810, 635], radius=18, fill=(61, 50, 28))
    draw.text((492, 602), "Quick action: Set weekly budget", fill=palette["paper"], font=font(16, True))
    image.save(SHOT_DIR / "03-after-phone.png")

    image, draw = canvas(
        "After - Fast Forward and Life Report reflection",
        "Fast Forward now presents time as a timeline, then hands off to a memory-style Life Report built from traces.",
        "04-after-fast-forward-report.png",
    )
    draw.rounded_rectangle([90, 150, 560, 640], radius=28, fill=(25, 25, 25), outline=palette["gold"], width=2)
    draw.text((126, 184), "Fast Forward", fill=palette["paper"], font=font(28, True))
    steps = ["Calendar pages move", "Costs and habits compound", "World pressure changes", "Life Report opens"]
    for i, step in enumerate(steps):
        y = 250 + i * 76
        draw.ellipse([132, y, 164, y + 32], fill=palette["gold"])
        draw.text((143, y + 7), str(i + 1), fill=palette["ink"], font=font(14, True))
        draw.text((184, y + 5), step, fill=palette["paper"], font=font(18, True))
    draw.rounded_rectangle([640, 150, 1190, 640], radius=28, fill=(25, 25, 25), outline=(80, 80, 76), width=2)
    draw.text((676, 184), "Life Report as Memory", fill=palette["paper"], font=font(28, True))
    chapters = ["Opening Reflection", "Major Moments", "Cause And Effect", "World Pressure", "Questions"]
    for i, chapter in enumerate(chapters):
        y = 250 + i * 64
        draw.text((680, y), f"{i + 1:02}", fill=palette["gold"], font=font(20, True))
        draw.text((732, y), chapter, fill=palette["paper"], font=font(18, True))
        draw.line([732, y + 34, 1130, y + 34], fill=(64, 64, 60), width=1)
    image.save(SHOT_DIR / "04-after-fast-forward-report.png")


def add_code_block(doc: Document, code: str):
    for line in code.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.style = "Code"
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(7)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc():
    DOCS.mkdir(parents=True, exist_ok=True)
    make_evidence_images()
    tests = run_tests()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 24, RGBColor(13, 13, 13)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
    if "Code" not in styles:
        styles.add_style("Code", 1)
    styles["Code"].font.name = "Courier New"
    styles["Code"].font.size = Pt(7)

    doc.add_heading("LifeVerse Review Volume 04 Implementation", 0)
    doc.add_paragraph("UX/UI implementation package for architect review. Source of truth priority: Volume 01, Volume 02, Volume 03, Volume 04.")
    doc.add_paragraph("Generated by Codex after implementing Volume 04 inside the existing LifeVerse / Life Sim layer.")

    doc.add_heading("1. Executive Summary", 1)
    add_bullets(doc, [
        "LifeVerse was refactored from tab/dashboard presentation toward a world-first game UX.",
        "Permanent HUD is reduced to time, day, money, health, energy, and current reminder.",
        "Phone, Journal, Map, and Pause now act as in-game device/menu overlays.",
        "Context interaction now shows location-relevant actions near the world view.",
        "Fast Forward and Life Report presentation now use timeline/chapter structures.",
        "Gameplay actions still route through Command Bus, Simulation Engine, Event Bus, Trace Engine, State Store, and UI.",
    ])

    doc.add_heading("2. Files Changed", 1)
    add_bullets(doc, CHANGED_FILES)
    doc.add_heading("3. New Files Added", 1)
    add_bullets(doc, NEW_FILES)

    doc.add_heading("4. Screenshots / Implementation Evidence", 1)
    doc.add_paragraph("Live screenshot capture was attempted but blocked by the local headless Chrome GPU failure and incomplete bundled Playwright package. To avoid fake evidence, the first image is explicitly labeled as a before reference reconstruction from the implementation diff, not a live screenshot. The after images are implementation evidence diagrams generated from the new Volume 04 UX structure.")
    for image_name, caption in [
        ("01-before-reference-dashboard.png", "Before reference: dashboard-like issue recorded from previous structure."),
        ("02-after-world-first.png", "After: world-first LifeVerse shell."),
        ("03-after-phone.png", "After: in-game phone overlay."),
        ("04-after-fast-forward-report.png", "After: Fast Forward and Life Report presentation."),
    ]:
        doc.add_paragraph(caption)
        doc.add_picture(str(SHOT_DIR / image_name), width=Inches(6.6))

    doc.add_heading("5. UX Changes By Volume 04 Chapter", 1)
    changes = [
        ("Chapter 03 Navigation", "Replaced traditional in-game tab bar with Phone, Journal, Map, and Pause dock. Query-param startup support was added for testing/review only."),
        ("Chapter 04 HUD", "Permanent HUD now exposes only critical information and moves detailed needs to dynamic warnings."),
        ("Chapter 05 Context Interaction", "Current location now exposes context actions such as Bed, Computer, Fridge, Workplace, Food Court, Gym Floor, Friend, and Quiet Bench."),
        ("Chapter 07 Smartphone UX", "Added phone overlay with Compass AI, Future Mirror, Calendar, Messages, Contacts, Maps, Banking, Journal, Tasks, and Settings."),
        ("Chapter 09 Life Report", "Report presentation uses reflective chapters: Opening Reflection, Major Moments, Cause And Effect, World Pressure, Lessons, Next Steps, Questions."),
        ("Chapter 11 Fast Forward", "Fast Forward panel now shows a time-lapse sequence before running simulation periods."),
        ("Chapter 13 Visual Design System", "Applied restrained black, warm off-white, gold accent, glass panels, consistent spacing, and reduced clutter."),
        ("Chapter 14 Motion", "Added lightweight panel transition and timeline animation-ready structure. Final art and motion polish remain Volume 05."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Volume 04 Area"
    table.rows[0].cells[1].text = "Implementation"
    for chapter, detail in changes:
        row = table.add_row().cells
        row[0].text = chapter
        row[1].text = detail

    doc.add_heading("6. Architecture Flow", 1)
    doc.add_paragraph("Gameplay action flow:")
    doc.add_paragraph("UI -> Command Bus -> Simulation Engine -> Gameplay Systems -> Event Bus -> Trace Engine -> State Store -> UI")
    doc.add_paragraph("Volume 04 UX helpers are read-only presentation logic. They select HUD labels, phone apps, context prompts, timeline labels, and report chapters from existing state.")

    doc.add_heading("7. Full Source Code Of New Files", 1)
    for relative in ["lifeverse-ux.js", "tests/lifeverse-volume04-ux.test.js"]:
        doc.add_heading(relative, 2)
        add_code_block(doc, read_file(relative))

    doc.add_heading("8. Exact Changed Functions From Modified Files", 1)
    app_source = read_file("app.js")
    for function_name in CHANGED_FUNCTIONS:
        doc.add_heading(f"app.js - {function_name}", 2)
        add_code_block(doc, extract_function(app_source, function_name))

    doc.add_heading("9. Styles Added For HUD / Phone / Reports / Fast Forward", 1)
    styles_source = read_file("styles.css")
    marker = "body.life-sim-active .volume04-world-first.lifeverse-shell"
    start = styles_source.find(marker)
    if start >= 0:
        end_marker = ".life-sim-asset-status"
        end = styles_source.find(end_marker, start)
        add_code_block(doc, styles_source[start:end])
    else:
        add_code_block(doc, "/* Volume 04 CSS block not found. */")

    doc.add_heading("10. Tests And Full Test Output", 1)
    add_code_block(doc, tests)

    doc.add_heading("11. Known Issues", 1)
    add_bullets(doc, [
        "Live screenshot capture could not complete in this environment because Chrome headless GPU failed and bundled Playwright lacked playwright-core.",
        "The 3D world still uses the existing WebGL/procedural Life Sim layer. Final art direction, character assets, production animation, and visual polish belong to Volume 05.",
        "The new UX shell is presentation-layer focused. It does not add new gameplay systems beyond existing Volume 02 systems.",
        "Phone apps that jump to outer Compass sections intentionally preserve the existing Compass platform rather than duplicating those systems inside LifeVerse.",
    ])

    doc.add_heading("12. What Still Belongs To Volume 05", 1)
    add_bullets(doc, [
        "Production anime character model, rig, and animation set.",
        "Final Singapore-inspired environment art, lighting, materials, and VFX.",
        "Final icon artwork, phone UI artwork, HUD polish, and motion effects.",
        "Audio, haptics, camera polish, weather visuals, and full cinematic Fast Forward sequence.",
    ])

    doc.add_heading("13. Next Recommended Step", 1)
    doc.add_paragraph("Review this package against Volume 04. After approval, deploy the current build to Vercel, then begin Volume 05 visual production without changing the Volume 04 UX rules.")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
